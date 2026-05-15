"""
DOCX renderer — generates Word documents for CV and cover letter.
Three template flavors per document type: modern, executive, technical.
"""

import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ── colour palettes per template ──────────────────────────────────────────────

_PALETTES: Dict[str, Dict[str, Any]] = {
    "modern": {
        "accent": RGBColor(0x00, 0x66, 0xCC),
        "heading_size": 20,
        "body_size": 11,
        "subheading_size": 12,
    },
    "executive": {
        "accent": RGBColor(0x1A, 0x1A, 0x2E),
        "heading_size": 22,
        "body_size": 11,
        "subheading_size": 13,
    },
    "technical": {
        "accent": RGBColor(0x2E, 0x7D, 0x32),
        "heading_size": 18,
        "body_size": 10,
        "subheading_size": 11,
    },
}

_TEMPLATE_MAP = {
    # CV
    "cv_modern.html": "modern",
    "cv_executive.html": "executive",
    "cv_technical.html": "technical",
    # Cover letter
    "cl_modern.html": "modern",
    "cl_executive.html": "executive",
    "cl_technical.html": "technical",
}


def _palette(template_name: str) -> Dict[str, Any]:
    flavor = _TEMPLATE_MAP.get(template_name, "modern")
    return _PALETTES[flavor]


# ── low-level helpers ──────────────────────────────────────────────────────────

def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin bottom-border paragraph as a visual divider."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_after = Pt(4)


def _heading(doc: Document, text: str, palette: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(palette["subheading_size"])
    _set_run_color(run, palette["accent"])
    _add_horizontal_rule(doc)


def _body(doc: Document, text: str, size: int, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _bullet(doc: Document, text: str, size: int) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.size = Pt(size)


# ── CV renderer ───────────────────────────────────────────────────────────────

def render_cv_docx(
    cv_content: Dict[str, Any],
    profile: Dict[str, Any],
    template_name: str,
    output_path: str,
) -> bool:
    """
    Render a CV to DOCX.

    Args:
        cv_content: Generated CV content (executive_summary, tailored_experience, etc.)
        profile: Candidate profile dict
        template_name: e.g. "cv_modern.html" — determines visual flavor
        output_path: Destination file path (.docx)

    Returns:
        True on success, False on error.
    """
    try:
        palette = _palette(template_name)
        doc = Document()

        # ── narrow margins ──
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        personal = profile.get("personal_info", {})
        name = personal.get("name", "")
        email = personal.get("email", "")
        phone = personal.get("phone", "")
        location = personal.get("location", "")
        linkedin = personal.get("linkedin", "")

        # ── Name ──
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(palette["heading_size"])
        _set_run_color(name_run, palette["accent"])

        # ── Contact line ──
        contact_parts = [p for p in [email, phone, location, linkedin] if p]
        contact_p = doc.add_paragraph("  |  ".join(contact_parts))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(8)
        for run in contact_p.runs:
            run.font.size = Pt(palette["body_size"])

        # ── Executive Summary ──
        summary = cv_content.get(
            "executive_summary", profile.get("summary", "")
        )
        if summary:
            _heading(doc, "Professional Summary", palette)
            _body(doc, summary, palette["body_size"])

        # ── Experience ──
        experience = cv_content.get(
            "tailored_experience", profile.get("experience", [])
        )
        if experience:
            _heading(doc, "Experience", palette)
            for entry in experience:
                title = entry.get("title", "")
                company = entry.get("company", "")
                start = entry.get("start", "")
                end = entry.get("end", "Present")
                achievements = entry.get("achievements", [])

                role_p = doc.add_paragraph()
                role_p.paragraph_format.space_before = Pt(6)
                role_p.paragraph_format.space_after = Pt(1)
                title_run = role_p.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(palette["body_size"] + 1)
                if company:
                    role_p.add_run(f"  —  {company}")
                    for r in role_p.runs[1:]:
                        r.font.size = Pt(palette["body_size"])

                date_p = doc.add_paragraph(f"{start} – {end}")
                date_p.paragraph_format.space_after = Pt(2)
                for run in date_p.runs:
                    run.font.size = Pt(palette["body_size"] - 1)
                    run.italic = True

                for ach in achievements:
                    _bullet(doc, ach, palette["body_size"])

        # ── Education ──
        education = profile.get("education", [])
        if education:
            _heading(doc, "Education", palette)
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", edu.get("end", ""))
                line = f"{degree}" + (f"  —  {institution}" if institution else "") + (f"  ({year})" if year else "")
                _body(doc, line, palette["body_size"], bold=False)

        # ── Skills ──
        skills = profile.get("skills", {})
        highlighted = cv_content.get("highlighted_skills", [])
        tech = skills.get("technical", [])
        soft = skills.get("soft", [])
        skill_list = highlighted if highlighted else tech + soft
        if skill_list:
            _heading(doc, "Skills", palette)
            _body(doc, "  ·  ".join(skill_list), palette["body_size"])

        # ── ATS Keywords ──
        ats = cv_content.get("ats_keywords", [])
        if ats:
            _heading(doc, "Key Competencies", palette)
            _body(doc, "  ·  ".join(ats), palette["body_size"])

        # ── Certifications ──
        certs = profile.get("certifications", [])
        if certs:
            _heading(doc, "Certifications", palette)
            for cert in certs:
                name_c = cert if isinstance(cert, str) else cert.get("name", str(cert))
                _body(doc, f"• {name_c}", palette["body_size"])

        doc.save(output_path)
        return True

    except Exception as exc:
        print(f"[docx_renderer] CV error: {exc}", file=sys.stderr)
        return False


# ── Cover Letter renderer ──────────────────────────────────────────────────────

def render_cover_letter_docx(
    cl_content: Dict[str, Any],
    profile: Dict[str, Any],
    template_name: str,
    output_path: str,
    date: str = "",
) -> bool:
    """
    Render a cover letter to DOCX.

    Args:
        cl_content: CoverLetter content dict (paragraphs list)
        profile: Candidate profile dict
        template_name: e.g. "cl_modern.html"
        output_path: Destination file path (.docx)
        date: Optional date string; defaults to today

    Returns:
        True on success, False on error.
    """
    try:
        palette = _palette(template_name)
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

        personal = profile.get("personal_info", {})
        name = personal.get("name", "")
        email = personal.get("email", "")
        phone = personal.get("phone", "")
        location = personal.get("location", "")

        if not date:
            date = datetime.now().strftime("%B %d, %Y")

        # ── Sender header ──
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(palette["heading_size"] - 2)
        _set_run_color(name_run, palette["accent"])
        name_p.paragraph_format.space_after = Pt(2)

        contact_parts = [p for p in [email, phone, location] if p]
        if contact_parts:
            contact_p = doc.add_paragraph("  |  ".join(contact_parts))
            contact_p.paragraph_format.space_after = Pt(2)
            for run in contact_p.runs:
                run.font.size = Pt(palette["body_size"])

        # Date
        date_p = doc.add_paragraph(date)
        date_p.paragraph_format.space_before = Pt(12)
        date_p.paragraph_format.space_after = Pt(12)
        for run in date_p.runs:
            run.font.size = Pt(palette["body_size"])

        _add_horizontal_rule(doc)

        # ── Paragraphs ──
        paragraphs = cl_content.get("paragraphs", [])
        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(para_text)
            run.font.size = Pt(palette["body_size"])

        doc.save(output_path)
        return True

    except Exception as exc:
        print(f"[docx_renderer] CoverLetter error: {exc}", file=sys.stderr)
        return False
